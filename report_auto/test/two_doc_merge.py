import os

from docx import Document
from docx.shared import Inches


def replace_string_with_image(doc, placeholder_text, image_path):
    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:
            # 清空段落内容
            paragraph.clear()
            # 在该段落中插入图片
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(1.0))  # 可以调整宽度等属性
            return True
    return False


# 定义文件路径
document_path = r'C:\Users\Administrator\Downloads\template\RBCD_L1_ECUMST_Test_Case_V03.docx'
image_path = r'E:\intellij_workspaces\report_auto\report_auto\static\images\dev_ops.png'
placeholder_text = '{10}'

# 检查文件是否存在
if not os.path.exists(document_path):
    print("The document does not exist.")
    exit()

# 检查图片文件是否存在
if not os.path.exists(image_path):
    print("The image file does not exist.")
    exit()

# 打开文档
doc = Document(document_path)

# 替换字符串并插入图片
if replace_string_with_image(doc, placeholder_text, image_path):
    # 保存文档
    doc.save(document_path)
    print(f"Image inserted at the placeholder '{placeholder_text}'.")
else:
    print(f"No placeholder '{placeholder_text}' found in the document.")
