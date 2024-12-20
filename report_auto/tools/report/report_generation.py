import logging
from typing import List

import pandas as pd
from pandas import DataFrame

from app.service.msttest.MstTestService import get_operator_rslt
from constant.faultType import FAULT_TYPE_MAPPING
from pojo.MSTReqPOJO import ReqPOJO
from tools.common.csv_column_rename import find_columns_with_dfc_err_type

pd.set_option('future.no_silent_downcasting', True)

import matplotlib.pyplot as plt
import matplotlib

matplotlib.use('Agg')  # TkAgg 或者 'Qt5Agg'

from docx import Document
from docx.shared import Inches

from tools.common.report_common import replace_placeholders_in_docx
from tools.common.dat_csv_common import create_file_path

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')


# 检查columns中每一列是否包含特定的fault_type值
def check_fault_type(df, columns, fault_type):
    # 创建一个空列表来存储包含故障类型的列
    columns_with_fault_type = []

    # 遍历指定的列
    for column in columns:
        # 检查该列是否包含给定的 fault_type
        if df[column].isin([fault_type]).any():
            columns_with_fault_type.append(column)

    return columns_with_fault_type


def replace_placeholders_in_boa_tables(doc, replacements):
    # 处理表格
    replace_placeholders_in_docx(doc, replacements)


'''
signals: 信号列
'''


def get_signals(fault_detection_df: DataFrame, docTemplateName: str, signals: list):
    # 异常类型
    fault_type: str = FAULT_TYPE_MAPPING.get(docTemplateName)
    # 异常类型数值列
    fault_type_column: str = f"DFC_st.{fault_type}"

    # 异常类型
    fault_type: str = FAULT_TYPE_MAPPING.get(docTemplateName)
    if fault_type is not None:
        columns_with_faults = find_columns_with_dfc_err_type(fault_detection_df, fault_type)
        if len(columns_with_faults) > 0:
            dfes_numdfc_column = columns_with_faults[0]

            # 异常类型对应数值
            fault_detection_df.loc[:, dfes_numdfc_column] = fault_detection_df[fault_type_column]

            # 绘制每个信号
            signals = signals + columns_with_faults
            logging.info(f"信号列:{signals}")

    return signals, fault_detection_df


'''
signals: 信号列
'''


def draw_img_in_boa_doc(fault_detection_df: pd.DataFrame, img_output_path: str, docTemplateName: str,
                        signals: List[str]) -> str:
    try:
        # 获取最终的信号列表和 DataFrame
        signals, fault_detection_df = get_signals(fault_detection_df, docTemplateName, signals)

        # 确保 'timestamps' 在信号列表中
        if 'timestamps' not in signals:
            signals.append('timestamps')
        fault_detection_df = fault_detection_df.loc[:, signals]

        # 找出每个信号的最大值
        max_values = fault_detection_df.drop(columns=['timestamps']).max()

        # 分类信号
        small_y_axis_signals = [col for col in max_values.index if max_values[col] <= 100]
        large_y_axis_signals = [col for col in max_values.index if max_values[col] > 100]

        # 创建一个新的图形
        fig, ax1 = plt.subplots(figsize=(10, 6))

        # 绘制第一个Y轴的信号
        for signal in small_y_axis_signals:
            ax1.plot(fault_detection_df['timestamps'], fault_detection_df[signal], label=signal)
        ax1.set_xlabel('Timestamps (s)')
        ax1.set_ylabel('Signal Values(L)', color='blue')
        ax1.set_ylim(0, 100)  # 设置第一个Y轴的范围从0到100
        ax1.yaxis.set_major_locator(plt.MultipleLocator(10))  # 设置第二个Y轴的刻度间隔为100
        ax1.tick_params(axis='y', labelcolor='blue')

        # 添加图例
        if small_y_axis_signals:
            ax1.legend(loc='upper left')  # 图例显示在左边

        # 如果有大于100的信号，创建第二个Y轴
        if large_y_axis_signals:
            max_large_value = max_values[large_y_axis_signals].max()
            ax2 = ax1.twinx()
            for signal in large_y_axis_signals:
                ax2.plot(fault_detection_df['timestamps'], fault_detection_df[signal], label=signal, linestyle='--')
            ax2.set_ylabel('Signal Values(R)', color='red')
            ax2.set_ylim(100, max_large_value + 100)  # 动态设置第二个Y轴的范围
            ax2.yaxis.set_major_locator(plt.MultipleLocator(100))  # 设置第二个Y轴的刻度间隔为100
            ax2.tick_params(axis='y', labelcolor='red')

            # 设置第二个Y轴的位置
            ax2.spines['right'].set_position(('outward', 60))  # 可以调整数值来改变位置
            ax2.yaxis.set_label_coords(1.15, 0.5)  # 调整Y轴标签的位置

            # 添加图例
            ax2.legend(loc='upper right')  # 图例显示在右边

        # 设置图表标题
        ax1.set_title('Signals Over Time')

        # 调整布局
        plt.tight_layout()

        # 保存图表到文件
        plt.savefig(img_output_path)

        # 关闭图形
        plt.close(fig)

        return img_output_path

    except Exception as e:
        print(f"An error occurred: {e}")
        return ""


"""
replacements:占位符及其对应值
fault_detection_df: 故障集
docTemplate: 模板名称
"""


def replace_variables_in_doc(replacements, fault_detection_df, signals: list, req_data: ReqPOJO) -> str:
    template_file_name, template_path = create_file_path(req_data.template_name, 'docx', req_data.template_path,
                                                         'template')
    img_name, img_path = create_file_path(req_data.template_name, 'png', req_data.output_path, 'img')
    doc_name, output_path = create_file_path(req_data.doc_output_name, 'docx', req_data.output_path, 'docx')

    draw_img_in_boa_doc(fault_detection_df, img_path, req_data.template_name, signals)
    logging.info(f"图片路径{img_path}")

    # 加载模板文档
    doc = Document(template_path)
    logging.info(f"模板路径:{template_path}")

    # 替换表格中的占位符
    logging.info(f"模板参数{replacements}")
    replace_placeholders_in_boa_tables(doc, replacements)

    # 插入图片
    doc.add_picture(img_path, width=Inches(6), height=Inches(4))

    # 保存输出文档
    doc.save(output_path)
    logging.info(f"文档路径:{output_path}")

    # 删除临时文件，释放磁盘空间
    # delete_file(req_data.csv_path)
    # delete_file(img_path)
    return output_path


def convert_value(value):
    return "√" if value == '1' else "❌" if value == '2' else ""


# mst首页报告生成结果
def replace_blank(header_file_path: str, clientIp: str, file_name: str) -> None:
    try:
        file_operator_rslt = get_operator_rslt(clientIp, file_name)
        logging.info("Query:%s",file_operator_rslt)
        # 定义要检查的键和结果变量的映射
        keys_to_check = {
            "app_pl_br_1": "file_operator_rslt_17",
            "brk_04": "file_operator_rslt_18",
            "brk_05": "file_operator_rslt_19",
            "ngs_06": "file_operator_rslt_20",
            "clth_05": "file_operator_rslt_21",
            "clth_06": "file_operator_rslt_22"
        }

        # 初始化结果变量并应用 convert_value 函数
        results = {result_var: convert_value(file_operator_rslt.get(key, "")) for key, result_var in
                   keys_to_check.items()}
        logging.info("ReplaceMent:%s",results)

        # 构建替换字典，只包括实际存在的键
        replacements = {
            f"{{{i}}}": results[result_var] for i, result_var in enumerate(keys_to_check.values(), start=17)
            if result_var in results
        }

        # 加载文档并替换占位符
        header_doc = Document(header_file_path)
        replace_placeholders_in_boa_tables(header_doc, replacements)

        # 保存更新后的文档
        header_doc.save(header_file_path)  # 或者保存到新的文件路径

    except Exception as e:
        print(f"An error occurred: {e}")
